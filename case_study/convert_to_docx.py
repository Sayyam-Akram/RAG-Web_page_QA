import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT_DIR = os.path.dirname(BASE_DIR)
CASE_STUDY_TXT = os.path.join(BASE_DIR, "case_study.txt")
IMAGES_DIR = os.path.join(ROOT_DIR, "case_study_images")
OUTPUT_DOCX = os.path.join(BASE_DIR, "CiteFlow_Case_Study.docx")

def add_markdown_paragraph(doc, text, style=None):
    """Parses simple markdown bold (**text**) and adds to doc."""
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    
    # Split by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p

def generate_docx():
    doc = Document()
    
    # Set Metadata
    doc.core_properties.title = "CiteFlow Case Study"
    doc.core_properties.author = "Sayyam Akram"

    # Read Content
    if not os.path.exists(CASE_STUDY_TXT):
        print(f"Error: {CASE_STUDY_TXT} not found.")
        return

    with open(CASE_STUDY_TXT, "r", encoding="utf-8") as f:
        lines = f.readlines()

    # Simple Markdown Parser
    in_code_block = False
    
    for line in lines:
        line = line.strip()
        
        if not line:
            continue
            
        # Code Blocks
        if line.startswith("```"):
            in_code_block = not in_code_block
            continue
        
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'Quote' # Use Quote style for code approximation
            continue

        # Headers
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        
        # Horizontal Rule
        elif line.startswith("---"):
            doc.add_page_break()
            
        # Lists
        elif line.startswith("- ") or line.startswith("* "):
            add_markdown_paragraph(doc, line[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            # Ordered list
            text = re.sub(r'^\d+\.\s', '', line)
            add_markdown_paragraph(doc, text, style='List Number')
            
        # Blockquotes
        elif line.startswith("> "):
            p = add_markdown_paragraph(doc, line[2:], style='Quote')
            
        # Tables (Skip complex parsing, just dump as text for now)
        elif "|" in line:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.style = 'No Spacing'

        # Normal Text
        else:
            # Skip image placeholders
            if line.startswith("!["):
                continue
            add_markdown_paragraph(doc, line)

    # Append Images
    doc.add_page_break()
    doc.add_heading("Appendix: Visual Evidence", level=1)
    
    if os.path.exists(IMAGES_DIR):
        image_files = sorted([f for f in os.listdir(IMAGES_DIR) if f.lower().endswith(('.png', '.jpg', '.jpeg'))])
        
        for img_file in image_files:
            img_path = os.path.join(IMAGES_DIR, img_file)
            try:
                doc.add_heading(f"Figure: {img_file}", level=3)
                doc.add_picture(img_path, width=Inches(6.0))
                doc.add_paragraph("") # Spacing
                print(f"Added image: {img_file}")
            except Exception as e:
                print(f"Failed to add image {img_file}: {e}")
    else:
        print(f"Images directory {IMAGES_DIR} not found.")

    # Save
    doc.save(OUTPUT_DOCX)
    print(f"Document saved to: {OUTPUT_DOCX}")

if __name__ == "__main__":
    generate_docx()
